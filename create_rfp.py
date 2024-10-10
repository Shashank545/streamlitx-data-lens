import streamlit as st
from datetime import date
import requests
import time
import docx
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import os, io
import base64

# Load configuration from config.json
import json

with open('./config_train.json', 'r') as config_file:
    config = json.load(config_file)

# Set up OpenAI API key
openai_api_key = config["openai_api_key"]

# Load prompts
intro_prompt = config["prompts"]["intro"]
scope_prompt = config["prompts"]["scope"]
deliverables_prompt = config["prompts"]["deliverables"]
vendor_guidelines_prompt = config["prompts"]["vendor_guidelines"]
training_kt_prompt = config["prompts"]["training_and_knowledge_transfer"]
index_prompt = config["prompts"]["index"]

intro_max_tokens = config["max_tokens"]["intro_max"]
scope_max_tokens = config["max_tokens"]["scope_max"]
deliverables_max_tokens = config["max_tokens"]["deliverables_max"]
vendor_max_tokens = config["max_tokens"]["vendor_inst_max"]
trainsupport_max_tokens = config["max_tokens"]["trainsupport_max"]
index_max_tokens = config["max_tokens"]["index_max"]

# Define Streamlit layout and widgets
st.title("RFP Document Generator")
st.markdown("---")

st.sidebar.header("User Instructions")
user_instructions = st.sidebar.text_area("Enter user instructions...", height=200)

st.sidebar.header("Intent")
intent = st.sidebar.selectbox("Select intent...", [
    'Queue Management System',
    'Cafeteria Enhancement',
    'Building Maintenace & Cleaning',
    'Advertising & Sales Promotion',
    'Laptop & IT hardware procurement',
    'Office gardening and beautification'
])

st.sidebar.header("Dates")
start_date = st.sidebar.date_input("Start Date", date.today())
end_date = st.sidebar.date_input("End Date", date.today())
rfp_ack_date = st.sidebar.date_input("RFP ACK Date", date.today())
rfp_submission_date = st.sidebar.date_input("RFP Submission Date", date.today())

st.sidebar.header("Contacts")
contact_emails = st.sidebar.multiselect("Select contact emails...", [
    'rfp_projects@damacgroup.com',
    'rfp_sales@damacgroup.com',
    'rfp_logistics@damacgroup.com',
    'rfp_backoffice@damacgroup.com',
    'rfp_genairealty@damacgroup.com',
    'rfp_finance@damacgroup.com',
    'rfp_hr@damacgroup.com'
])

# Function to generate content using OpenAI API
def generate_section_content(prompt, section_max_token, intent):
    api_url = "https://api.openai.com/v1/completions"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {openai_api_key}"
    }
    payload = {
        "model": "gpt-3.5-turbo-instruct",
        "max_tokens": int(section_max_token),
        "prompt": f"{prompt}\nRFP Intent: {intent}",
        "temperature": 0.1,
        "n": 1
    }

    # Retry with backoff strategy
    max_retries = 5
    for retry in range(max_retries):
        response = requests.post(api_url, json=payload, headers=headers)
        if response.status_code == 429:
            st.warning(f"Rate limited, retrying in {2**retry} seconds...")
            time.sleep(2**retry)
        else:
            break

    return response.json()["choices"][0]["text"]

# Generate RFP content on button click
if st.button("SUBMIT"):
    st.markdown("---")
    st.info("Generating RFP content...")

    introduction = generate_section_content(f"{intro_prompt} {user_instructions}", intro_max_tokens, intent)
    scope_of_rfp = generate_section_content(f"{scope_prompt} {user_instructions}", scope_max_tokens, intent)
    deliverables_of_rfp = generate_section_content(f"{deliverables_prompt} {user_instructions}", deliverables_max_tokens, intent)
    vendor_instructions = generate_section_content(f"{vendor_guidelines_prompt} {user_instructions} {start_date} {end_date} {rfp_ack_date} {rfp_submission_date} {contact_emails}", vendor_max_tokens, intent)
    trainingsupport = generate_section_content(f"{training_kt_prompt}", trainsupport_max_tokens, intent)
    index_page = generate_section_content(f"{index_prompt}", index_max_tokens, intent)

    st.success("RFP content generated successfully!")

    # Display generated content (optional)
    st.markdown("---")
    st.header("Generated RFP Content")
    st.subheader("Introduction")
    st.write(introduction)
    st.subheader("Scope of RFP")
    st.write(scope_of_rfp)
    st.subheader("Deliverables of RFP")
    st.write(deliverables_of_rfp)
    st.subheader("Vendor Instructions")
    st.write(vendor_instructions)
    st.subheader("Training and Knowledge Transfer")
    st.write(trainingsupport)
    st.subheader("Index Page")
    st.write(index_page)

    # Enable download buttons
    st.markdown("---")
    st.markdown("### Download Options")
    st.success("RFP content is ready for download.")

    # Example: Download as DOCX
    docx_filename = "rfp_document.docx"
    doc = Document()
    
    # Add cover page
    # doc.add_heading('Request for Proposal (RFP) Document', level=1).bold = True
    # doc.add_paragraph(f"Intent: {intent}")
    # doc.add_paragraph(f"User Instructions: {user_instructions}")
    # doc.add_paragraph(f"Start Date of RFP: {start_date}")
    # doc.add_paragraph(f"End Date of RFP: {end_date}")
    # doc.add_paragraph(f"RFP ACK Date: {rfp_ack_date}")
    # doc.add_paragraph(f"RFP Submission Date: {rfp_submission_date}")
    # doc.add_paragraph(f"Contacts: {', '.join(contact_emails)}")
    # doc.add_page_break()
    # Cover Page Title
    # Insert DAMAC Logo at the top left with reduced size
    image_path = 'damac_logo.jpg'
    width = Inches(2)  # Adjust width as needed
    height = Inches(0.25)  # Adjust height as needed
    logo_paragraph = doc.add_paragraph()
    logo_run = logo_paragraph.add_run()
    logo_run.add_picture(image_path, width=width, height=height)

    # Optionally, you can align the logo to the left
    logo_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    doc.add_heading('Request for Proposal (RFP) Document', level=1).bold = True
    # image_path = 'damac_logo.jpg'
    # doc.add_picture(image_path, width=docx.shared.Inches(6))
    # doc.add_page_break()

    # Data for the cover page
    data = {
        "Intent": intent,
        "Start Date of RFP": start_date,
        "End Date of RFP": end_date,
        "RFP ACK Date": rfp_ack_date,
        "RFP Submission Date": rfp_submission_date,
        "Contacts": ', '.join(contact_emails)
    }

    # Add table for cover page
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'

    # Populate table with data
    for field_name, field_value in data.items():
        row_cells = table.add_row().cells
        row_cells[0].text = field_name
        row_cells[1].text = str(field_value)

    # Add page break after cover page
    doc.add_page_break()

    # Add index page
    doc.add_heading('Table of Contents', level=1).bold = True
    index_content = index_page.split('\n')
    for line in index_content:
        doc.add_paragraph(line, style='BodyText')
    doc.add_page_break()

    # Add sections
    doc.add_heading('Introduction', level=1).bold = True
    doc.add_paragraph(introduction, style='BodyText')
    doc.add_page_break()

    # Systems and Applications Section
    systems_and_applications_content = (
        "DAMAC is undergoing multiple transformation programs which also entail technology implementation. "
        "Oracle Fusion Financials is the ERP system being utilized since Oct-2022 for most DAMAC Entities "
        "which includes DAMAC Properties, LOAMS, Data Center, Leasing, DAMAC Capital. However, there are certain "
        "group entities whose ERP systems do not utilize Oracle Fusion as their ERP system."
    )

    doc.add_heading('Systems and Applications', level=1)
    doc.add_paragraph(systems_and_applications_content, style='BodyText')

    image_path = 'rfp_image_01.png'
    doc.add_picture(image_path, width=docx.shared.Inches(6))
    doc.add_page_break()

    doc.add_heading('Scope of RFP', level=1).bold = True
    doc.add_paragraph(scope_of_rfp, style='BodyText')
    doc.add_page_break()

    doc.add_heading('Deliverables of RFP', level=1).bold = True
    doc.add_paragraph(deliverables_of_rfp, style='BodyText')
    doc.add_page_break()

    doc.add_heading('Vendor Instructions', level=1).bold = True
    doc.add_paragraph(vendor_instructions, style='BodyText')
    doc.add_page_break()

    doc.add_heading('Training and Knowledge Transfer', level=1).bold = True
    doc.add_paragraph(trainingsupport, style='BodyText')
    doc.add_page_break()

    # Save DOCX file
    with io.BytesIO() as output:
        doc.save(output)
        docx_data = output.getvalue()

    st.markdown("---")
    st.markdown("### Download RFP Document (DOCX)")
    st.markdown(f"Download [rfp_document.docx](data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{base64.b64encode(docx_data).decode()})")

